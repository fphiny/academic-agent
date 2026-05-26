from __future__ import annotations

import os
import re
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import unquote, urlparse

import requests
from bs4 import BeautifulSoup, NavigableString
from pypdf import PdfReader

from ..base import (
    ToolResult,
    clean_block_text,
    normalize_whitespace,
)

# -------------------------
# limits
# -------------------------
MAX_DOC_CHARS = 200_000
MAX_BLOCK_CHARS = 12_000
DEFAULT_TIMEOUT = 20

DEFAULT_REQUEST_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/122.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7",
    "Cache-Control": "no-cache",
    "Pragma": "no-cache",
    "Connection": "keep-alive",
}

HTML_MIME_TYPES = {
    "text/html",
    "application/xhtml+xml",
}

TEXT_MIME_TYPES = {
    "text/plain",
    "text/markdown",
    "application/json",
    "text/csv",
    "application/xml",
    "text/xml",
}

PDF_MIME_TYPES = {
    "application/pdf",
}

DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

GENERIC_BINARY_MIME_TYPES = {
    "application/octet-stream",
    "application/x-msdownload",
    "application/download",
    "application/x-download",
    "binary/octet-stream",
}

HWP_MIME_TYPES = {
    "application/x-hwp",
    "application/haansofthwp",
    "application/hwp",
    "application/octet-stream",
    "application/x-msdownload",
}


@dataclass
class DocumentBlock:
    url: str
    title: str
    content_type: str
    block_id: str
    heading: str
    text: str
    page: Optional[int] = None


def safe_truncate(text: str, max_chars: int) -> str:
    text = normalize_whitespace(text or "")
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + " ..."


def _base_content_type(content_type: str) -> str:
    return (content_type or "").split(";")[0].strip().lower()


def _extract_filename_from_url(url: str) -> str:
    try:
        path = unquote(urlparse(url).path)
        name = Path(path).name
        return name or ""
    except Exception:
        return ""


def _extract_filename_from_cd(content_disposition: str) -> str:
    """
    Content-Disposition 에서 실제 파일명 추출
    예:
    attachment; filename="abc.hwp"
    attachment; filename*=UTF-8''%ED%95%9C%EA%B8%80.hwp
    """
    if not content_disposition:
        return ""

    cd = content_disposition.strip()

    m = re.search(r"filename\*\s*=\s*UTF-8''([^;]+)", cd, re.I)
    if m:
        return unquote(m.group(1)).strip().strip('"')

    m = re.search(r'filename\s*=\s*"([^"]+)"', cd, re.I)
    if m:
        return m.group(1).strip()

    m = re.search(r"filename\s*=\s*([^;]+)", cd, re.I)
    if m:
        return m.group(1).strip().strip('"')

    return ""


def _file_suffix(filename: str) -> str:
    try:
        return Path(filename).suffix.lower()
    except Exception:
        return ""


def _resolved_filename(url: str, final_url: str, content_disposition: str) -> str:
    return (
        _extract_filename_from_cd(content_disposition)
        or _extract_filename_from_url(final_url)
        or _extract_filename_from_url(url)
    )


def _default_title(
    url: str,
    final_url: str = "",
    content_disposition: str = "",
    fallback: str = "",
) -> str:
    resolved_name = _resolved_filename(url, final_url, content_disposition)
    if resolved_name:
        return clean_block_text(resolved_name)
    return clean_block_text(fallback or "")


def _is_pdf_bytes(data: bytes) -> bool:
    return data.startswith(b"%PDF")


def _is_zip_bytes(data: bytes) -> bool:
    return data.startswith(b"PK\x03\x04")


def _is_ole_bytes(data: bytes) -> bool:
    return data.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1")


def _looks_like_html_bytes(data: bytes) -> bool:
    if not data:
        return False

    head = data[:4096].lstrip()
    if not head:
        return False

    if b"\x00" in head[:512]:
        return False

    try:
        sample = head.decode("utf-8", errors="ignore").lower()
    except Exception:
        try:
            sample = head.decode("latin-1", errors="ignore").lower()
        except Exception:
            return False

    sample = sample.strip()
    if not sample:
        return False

    html_markers = (
        "<!doctype html",
        "<html",
        "<head",
        "<body",
        "<meta",
        "<title",
        "<script",
    )

    if any(marker in sample for marker in html_markers):
        return True

    if sample.startswith("<") and re.search(r"<[a-z][a-z0-9:_-]*", sample):
        return True

    return False


def _looks_like_text_bytes(data: bytes) -> bool:
    if not data:
        return False

    head = data[:4096]
    if not head:
        return False

    if b"\x00" in head:
        return False

    try:
        decoded = head.decode("utf-8", errors="ignore")
    except Exception:
        try:
            decoded = head.decode("latin-1", errors="ignore")
        except Exception:
            return False

    printable = sum(1 for ch in decoded if ch.isprintable() or ch in "\r\n\t")
    total = max(len(decoded), 1)
    return (printable / total) >= 0.9


def _head_hex(data: bytes, limit: int = 16) -> str:
    try:
        return data[:limit].hex()
    except Exception:
        return ""


def _detect_zip_kind(data: bytes, filename: str = "") -> str:
    """
    DOCX/HWPX 등 ZIP 기반 포맷 판별
    """
    if not _is_zip_bytes(data):
        return "unknown"

    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            names = set(zf.namelist())

            if "[Content_Types].xml" in names and "word/document.xml" in names:
                return "docx"

            hwpx_markers = {
                "Contents/section0.xml",
                "Contents/content.hpf",
                "mimetype",
            }
            if any(name in names for name in hwpx_markers):
                if _file_suffix(filename) == ".hwpx":
                    return "hwpx"

            return "zip"
    except Exception:
        return "zip"


def _detect_ole_kind(data: bytes, filename: str = "") -> str:
    """
    OLE Compound File 안에서 진짜 HWP인지 확인
    """
    if not _is_ole_bytes(data):
        return "unknown"

    suffix = _file_suffix(filename)

    try:
        import olefile
    except Exception:
        if suffix == ".hwp":
            return "hwp"
        if suffix == ".doc":
            return "doc"
        return "ole"

    try:
        with olefile.OleFileIO(BytesIO(data)) as ole:
            stream_names = set("/".join(x) for x in ole.listdir())

            if "FileHeader" in stream_names:
                try:
                    raw = ole.openstream("FileHeader").read()
                    if b"HWP Document File" in raw:
                        return "hwp"
                except Exception:
                    pass

                return "hwp"

            if suffix == ".doc":
                return "doc"
            if suffix == ".hwp":
                return "hwp"

            return "ole"
    except Exception:
        if suffix == ".hwp":
            return "hwp"
        if suffix == ".doc":
            return "doc"
        return "ole"


def _guess_kind(
    *,
    url: str,
    final_url: str,
    content_type: str,
    content_disposition: str,
    content: bytes,
) -> str:
    """
    판별 우선순위:
    1) magic bytes
    2) Content-Disposition / filename suffix
    3) MIME
    4) content sniffing
    5) URL suffix
    """
    ct = _base_content_type(content_type)
    preferred_filename = _resolved_filename(url, final_url, content_disposition)
    suffix = _file_suffix(preferred_filename)

    # 1) magic bytes 최우선
    if _is_pdf_bytes(content):
        return "pdf"

    zip_kind = _detect_zip_kind(content, preferred_filename)
    if zip_kind == "docx":
        return "docx"
    if zip_kind == "hwpx":
        return "hwpx"

    ole_kind = _detect_ole_kind(content, preferred_filename)
    if ole_kind == "hwp":
        return "hwp"

    # 2) filename suffix 우선
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".hwp":
        return "hwp"
    if suffix == ".hwpx":
        return "hwpx"
    if suffix in {".txt", ".md", ".csv", ".json", ".xml", ".log"}:
        return "text"
    if suffix in {".html", ".htm"}:
        return "html"

    # 3) MIME
    if ct in PDF_MIME_TYPES:
        return "pdf"
    if ct in DOCX_MIME_TYPES:
        return "docx"
    if ct in HWP_MIME_TYPES:
        return "hwp"
    if ct in TEXT_MIME_TYPES:
        return "html" if _looks_like_html_bytes(content) else "text"
    if ct in HTML_MIME_TYPES:
        return "html"

    # 4) content sniffing
    if _looks_like_html_bytes(content):
        return "html"
    if _looks_like_text_bytes(content):
        return "text"

    # 5) generic binary fallback
    if ct in GENERIC_BINARY_MIME_TYPES:
        return "unknown"

    # 6) URL suffix 마지막
    url_suffix = _file_suffix(_extract_filename_from_url(final_url) or _extract_filename_from_url(url))
    if url_suffix == ".pdf":
        return "pdf"
    if url_suffix == ".docx":
        return "docx"
    if url_suffix == ".hwp":
        return "hwp"
    if url_suffix == ".hwpx":
        return "hwpx"
    if url_suffix in {".txt", ".md", ".csv", ".json", ".xml", ".log"}:
        return "text"
    if url_suffix in {".html", ".htm"}:
        return "html"

    return "unknown"


def html_to_markdown(soup: BeautifulSoup, title: str = "") -> str:
    for img in soup.find_all("img"):
        alt = normalize_whitespace(img.get("alt", ""))
        if alt:
            img.replace_with(NavigableString(f" [image: {alt}] "))
        else:
            img.decompose()

    for tag in soup(["script", "style", "noscript", "svg", "iframe", "canvas"]):
        tag.decompose()

    body = soup.body or soup

    try:
        from markdownify import markdownify as md

        raw_html = str(body)
        markdown_text = md(
            raw_html,
            heading_style="ATX",
            bullets="-",
            strip=["script", "style", "noscript", "svg", "iframe", "canvas"],
        )
        markdown_text = normalize_whitespace(markdown_text)

        if title and not markdown_text.startswith("# "):
            markdown_text = f"# {clean_block_text(title)}\n\n{markdown_text}"

        return markdown_text

    except Exception:
        text = body.get_text("\n", strip=True)
        text = normalize_whitespace(text)
        if title:
            return f"# {clean_block_text(title)}\n\n{text}"
        return text


def extract_html_blocks(markdown_text: str, url: str, title: str) -> List[DocumentBlock]:
    normalized = normalize_whitespace(markdown_text)
    if not normalized:
        return []

    section_pattern = r"(?ms)(^#{1,6}\s+.*?)(?=^#{1,6}\s+|\Z)"
    matches = re.findall(section_pattern, normalized)

    if not matches:
        return [
            DocumentBlock(
                url=url,
                title=title,
                content_type="html",
                block_id="html-1",
                heading=title or "document",
                text=safe_truncate(normalized, MAX_BLOCK_CHARS),
            )
        ]

    blocks: List[DocumentBlock] = []
    for i, section in enumerate(matches, start=1):
        cleaned = normalize_whitespace(section)
        if not cleaned:
            continue

        first_line = cleaned.splitlines()[0].strip()
        heading = re.sub(r"^#{1,6}\s*", "", first_line).strip() or f"section-{i}"

        blocks.append(
            DocumentBlock(
                url=url,
                title=title,
                content_type="html",
                block_id=f"html-{i}",
                heading=heading,
                text=safe_truncate(cleaned, MAX_BLOCK_CHARS),
            )
        )
    return blocks


def extract_pdf_blocks_from_reader(
    reader: PdfReader,
    url: str,
    title: str = "",
) -> List[DocumentBlock]:
    blocks: List[DocumentBlock] = []

    for page_idx, page in enumerate(reader.pages, start=1):
        extracted = normalize_whitespace(page.extract_text() or "")
        if not extracted:
            continue

        blocks.append(
            DocumentBlock(
                url=url,
                title=title,
                content_type="pdf",
                block_id=f"pdf-page-{page_idx}",
                heading=f"Page {page_idx}",
                text=safe_truncate(extracted, MAX_BLOCK_CHARS),
                page=page_idx,
            )
        )

    return blocks


def extract_text_blocks(
    text: str,
    url: str,
    title: str = "",
    content_type: str = "text",
) -> List[DocumentBlock]:
    raw_text = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    raw_parts = re.split(r"\n\s*\n+", raw_text)
    blocks: List[DocumentBlock] = []

    for i, part in enumerate(raw_parts, start=1):
        cleaned = normalize_whitespace(part)
        if not cleaned:
            continue

        blocks.append(
            DocumentBlock(
                url=url,
                title=title,
                content_type=content_type,
                block_id=f"{content_type}-{i}",
                heading=f"Block {i}",
                text=safe_truncate(cleaned, MAX_BLOCK_CHARS),
            )
        )

    return blocks


def extract_docx_blocks_from_bytes(
    data: bytes,
    url: str,
    title: str = "",
) -> List[DocumentBlock]:
    try:
        from docx import Document as DocxDocument
    except Exception as e:
        raise RuntimeError("DOCX 지원을 위해 `python-docx` 설치가 필요합니다.") from e

    document = DocxDocument(BytesIO(data))

    doc_title = normalize_whitespace(getattr(document.core_properties, "title", "") or "")
    resolved_title = title or doc_title or _default_title(url)

    blocks: List[DocumentBlock] = []
    current_heading = resolved_title or "document"
    current_lines: List[str] = []
    section_index = 1

    def flush_section() -> None:
        nonlocal section_index, current_lines, current_heading
        text = normalize_whitespace("\n".join(current_lines))
        if not text:
            current_lines = []
            return

        blocks.append(
            DocumentBlock(
                url=url,
                title=resolved_title,
                content_type="docx",
                block_id=f"docx-section-{section_index}",
                heading=current_heading or f"Section {section_index}",
                text=safe_truncate(text, MAX_BLOCK_CHARS),
            )
        )
        section_index += 1
        current_lines = []

    for para in document.paragraphs:
        text = normalize_whitespace(para.text or "")
        if not text:
            continue

        style_name = normalize_whitespace(getattr(getattr(para, "style", None), "name", "") or "")
        if style_name.lower().startswith("heading"):
            flush_section()
            current_heading = text
            continue

        current_lines.append(text)

    flush_section()

    for table_idx, table in enumerate(document.tables, start=1):
        row_lines: List[str] = []
        for row in table.rows:
            cells = [normalize_whitespace(cell.text or "") for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                row_lines.append(" | ".join(cells))

        table_text = "\n".join(row_lines).strip()
        if not table_text:
            continue

        blocks.append(
            DocumentBlock(
                url=url,
                title=resolved_title,
                content_type="docx",
                block_id=f"docx-table-{table_idx}",
                heading=f"Table {table_idx}",
                text=safe_truncate(table_text, MAX_BLOCK_CHARS),
            )
        )

    if not blocks:
        full_text = normalize_whitespace(
            "\n".join(p.text for p in document.paragraphs if (p.text or "").strip())
        )
        if full_text:
            blocks.append(
                DocumentBlock(
                    url=url,
                    title=resolved_title,
                    content_type="docx",
                    block_id="docx-1",
                    heading=resolved_title or "document",
                    text=safe_truncate(full_text, MAX_BLOCK_CHARS),
                )
            )

    return blocks


def extract_hwp_blocks_from_bytes(
    data: bytes,
    url: str,
    title: str = "",
) -> List[DocumentBlock]:
    try:
        from langchain_teddynote.document_loaders import HWPLoader
    except Exception as e:
        raise RuntimeError(
            "HWP 지원을 위해 `langchain-teddynote` 설치가 필요합니다."
        ) from e

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as tmp:
            tmp.write(data)
            tmp.flush()
            tmp_path = tmp.name

        loader = HWPLoader(tmp_path)
        docs = loader.load()

        texts: List[str] = []
        for doc in docs:
            page_content = normalize_whitespace(getattr(doc, "page_content", "") or "")
            if page_content:
                texts.append(page_content)

        merged_text = "\n\n".join(texts).strip()
        if not merged_text:
            return []

        resolved_title = title or _default_title(url)
        return extract_text_blocks(
            merged_text,
            url=url,
            title=resolved_title,
            content_type="hwp",
        )

    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def _normalize_input_urls(
    urls: Optional[List[str]] = None,
    items: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    normalized: List[Dict[str, Any]] = []
    seen = set()

    if isinstance(items, list):
        for item in items:
            if not isinstance(item, dict):
                continue
            link = str(item.get("link") or item.get("url") or "").strip()
            if not link or link in seen:
                continue
            seen.add(link)
            normalized.append(
                {
                    "url": link,
                    "search_title": normalize_whitespace(str(item.get("title") or "")),
                    "search_snippet": normalize_whitespace(str(item.get("snippet") or "")),
                    "display_link": normalize_whitespace(str(item.get("displayLink") or "")),
                    "source_engine": normalize_whitespace(str(item.get("source_engine") or "")),
                }
            )

    if isinstance(urls, list):
        for url in urls:
            link = str(url or "").strip()
            if not link or link in seen:
                continue
            seen.add(link)
            normalized.append(
                {
                    "url": link,
                    "search_title": "",
                    "search_snippet": "",
                    "display_link": "",
                    "source_engine": "",
                }
            )

    return normalized


def _decode_response_text(resp: requests.Response) -> str:
    try:
        return resp.text
    except Exception:
        try:
            encoding = resp.encoding or resp.apparent_encoding or "utf-8"
            return resp.content.decode(encoding, errors="replace")
        except Exception:
            return resp.content.decode("utf-8", errors="replace")


def _build_headers_for_url(url: str) -> Dict[str, str]:
    headers = dict(DEFAULT_REQUEST_HEADERS)
    try:
        parsed = urlparse(url)
        if parsed.scheme and parsed.netloc:
            headers["Referer"] = f"{parsed.scheme}://{parsed.netloc}/"
    except Exception:
        pass
    return headers


def _log_final_document(debug_print, result: Dict[str, Any]) -> None:
    if not debug_print:
        return

    debug_print(
        "FETCH FINAL DOCUMENT",
        {
            "url": result.get("url"),
            "title": result.get("title"),
            "content_type": result.get("content_type"),
            "source_content_type": result.get("source_content_type"),
            "content_disposition": result.get("content_disposition"),
            "byte_size": result.get("byte_size"),
            "text_len": len(result.get("text", "") or ""),
            "raw_text_len": len(result.get("raw_text", "") or ""),
            "block_count": len(result.get("blocks", []) or []),
            "warning": result.get("warning", ""),
        },
    )


def _empty_or_shell_html(html_text: str) -> bool:
    if not html_text or not html_text.strip():
        return True

    trimmed = html_text.strip()
    if len(trimmed) < 30:
        return True

    try:
        soup = BeautifulSoup(html_text, "html.parser")
        body_text = normalize_whitespace((soup.body or soup).get_text(" ", strip=True))
    except Exception:
        body_text = normalize_whitespace(html_text)

    if body_text:
        return False

    lowered = trimmed.lower()
    if any(tag in lowered for tag in ("<html", "<body", "<head", "<script", "<meta", "<title")):
        return True

    return len(trimmed) < 200


def fetch_one_url(
    *,
    url: str,
    search_title: str = "",
    search_snippet: str = "",
    display_link: str = "",
    source_engine: str = "",
    timeout: int = DEFAULT_TIMEOUT,
    debug_print=None,
) -> Dict[str, Any]:
    try:
        with requests.Session() as session:
            resp = session.get(
                url,
                timeout=timeout,
                headers=_build_headers_for_url(url),
                allow_redirects=True,
            )

        resp.raise_for_status()

        content_type = resp.headers.get("content-type", "").lower()
        content_disposition = resp.headers.get("content-disposition", "")
        byte_size = len(resp.content or b"")

        detected_kind = _guess_kind(
            url=url,
            final_url=resp.url,
            content_type=content_type,
            content_disposition=content_disposition,
            content=resp.content,
        )

        resolved_title = _default_title(
            url=url,
            final_url=resp.url,
            content_disposition=content_disposition,
            fallback=search_title or url,
        )

        if debug_print:
            debug_print(
                f"EXTERNAL_FETCH RESPONSE META | {url}",
                {
                    "status_code": resp.status_code,
                    "content_type": content_type,
                    "content_disposition": content_disposition,
                    "final_url": resp.url,
                    "detected_kind": detected_kind,
                    "resolved_title": resolved_title,
                    "byte_size": byte_size,
                    "head_hex": _head_hex(resp.content),
                    "looks_like_html_bytes": _looks_like_html_bytes(resp.content),
                    "looks_like_text_bytes": _looks_like_text_bytes(resp.content),
                },
            )

        if detected_kind == "html":
            html_text = _decode_response_text(resp)

            if _empty_or_shell_html(html_text):
                result = {
                    "url": url,
                    "final_url": resp.url,
                    "title": resolved_title,
                    "content_type": "unknown",
                    "source_content_type": content_type,
                    "content_disposition": content_disposition,
                    "raw_html": safe_truncate(html_text, MAX_DOC_CHARS),
                    "raw_text": "",
                    "text": "",
                    "blocks": [],
                    "search_title": search_title,
                    "search_snippet": search_snippet,
                    "display_link": display_link,
                    "source_engine": source_engine,
                    "status_code": resp.status_code,
                    "byte_size": byte_size,
                    "warning": (
                        "html response is empty/shell-like; "
                        "likely blocked download endpoint, permission page, or attachment indirection"
                    ),
                }
                _log_final_document(debug_print, result)
                return result

            soup = BeautifulSoup(html_text, "html.parser")

            title = resolved_title
            if soup.title and soup.title.string:
                title = clean_block_text(soup.title.string) or resolved_title

            markdown_text = html_to_markdown(soup=soup, title=title)
            markdown_text = safe_truncate(markdown_text, MAX_DOC_CHARS)
            blocks = extract_html_blocks(markdown_text, url=url, title=title)

            result = {
                "url": url,
                "final_url": resp.url,
                "title": title,
                "content_type": "html",
                "source_content_type": content_type,
                "content_disposition": content_disposition,
                "raw_html": safe_truncate(html_text, MAX_DOC_CHARS),
                "raw_text": "",
                "text": markdown_text,
                "blocks": [asdict(b) for b in blocks],
                "search_title": search_title,
                "search_snippet": search_snippet,
                "display_link": display_link,
                "source_engine": source_engine,
                "status_code": resp.status_code,
                "byte_size": byte_size,
            }
            _log_final_document(debug_print, result)
            return result

        if detected_kind == "text":
            decoded_text = _decode_response_text(resp)
            raw_text = safe_truncate(decoded_text, MAX_DOC_CHARS)
            blocks = extract_text_blocks(raw_text, url=url, title=resolved_title, content_type="text")
            joined_text = safe_truncate("\n\n".join(block.text for block in blocks), MAX_DOC_CHARS)

            result = {
                "url": url,
                "final_url": resp.url,
                "title": resolved_title,
                "content_type": "text",
                "source_content_type": content_type,
                "content_disposition": content_disposition,
                "raw_html": "",
                "raw_text": raw_text,
                "text": joined_text,
                "blocks": [asdict(b) for b in blocks],
                "search_title": search_title,
                "search_snippet": search_snippet,
                "display_link": display_link,
                "source_engine": source_engine,
                "status_code": resp.status_code,
                "byte_size": byte_size,
            }
            _log_final_document(debug_print, result)
            return result

        if detected_kind == "pdf":
            reader = PdfReader(BytesIO(resp.content))
            blocks = extract_pdf_blocks_from_reader(reader, url=url, title=resolved_title)
            text = safe_truncate("\n".join(block.text for block in blocks), MAX_DOC_CHARS)

            result = {
                "url": url,
                "final_url": resp.url,
                "title": resolved_title,
                "content_type": "pdf",
                "source_content_type": content_type,
                "content_disposition": content_disposition,
                "raw_html": "",
                "raw_text": text,
                "text": text,
                "blocks": [asdict(b) for b in blocks],
                "search_title": search_title,
                "search_snippet": search_snippet,
                "display_link": display_link,
                "source_engine": source_engine,
                "status_code": resp.status_code,
                "byte_size": byte_size,
            }
            _log_final_document(debug_print, result)
            return result

        if detected_kind == "docx":
            blocks = extract_docx_blocks_from_bytes(resp.content, url=url, title=resolved_title)
            text = safe_truncate("\n\n".join(block.text for block in blocks), MAX_DOC_CHARS)

            result = {
                "url": url,
                "final_url": resp.url,
                "title": resolved_title,
                "content_type": "docx",
                "source_content_type": content_type,
                "content_disposition": content_disposition,
                "raw_html": "",
                "raw_text": text,
                "text": text,
                "blocks": [asdict(b) for b in blocks],
                "search_title": search_title,
                "search_snippet": search_snippet,
                "display_link": display_link,
                "source_engine": source_engine,
                "status_code": resp.status_code,
                "byte_size": byte_size,
            }
            _log_final_document(debug_print, result)
            return result

        if detected_kind == "hwp":
            try:
                blocks = extract_hwp_blocks_from_bytes(resp.content, url=url, title=resolved_title)
                text = safe_truncate("\n\n".join(block.text for block in blocks), MAX_DOC_CHARS)

                result = {
                    "url": url,
                    "final_url": resp.url,
                    "title": resolved_title,
                    "content_type": "hwp",
                    "source_content_type": content_type,
                    "content_disposition": content_disposition,
                    "raw_html": "",
                    "raw_text": text,
                    "text": text,
                    "blocks": [asdict(b) for b in blocks],
                    "search_title": search_title,
                    "search_snippet": search_snippet,
                    "display_link": display_link,
                    "source_engine": source_engine,
                    "status_code": resp.status_code,
                    "byte_size": byte_size,
                }
                _log_final_document(debug_print, result)
                return result

            except Exception as e:
                result = {
                    "url": url,
                    "final_url": resp.url,
                    "title": resolved_title,
                    "content_type": "hwp",
                    "source_content_type": content_type,
                    "content_disposition": content_disposition,
                    "raw_html": "",
                    "raw_text": "",
                    "text": "",
                    "blocks": [],
                    "search_title": search_title,
                    "search_snippet": search_snippet,
                    "display_link": display_link,
                    "source_engine": source_engine,
                    "status_code": resp.status_code,
                    "byte_size": byte_size,
                    "warning": f"hwp extraction failed: {e}",
                }
                _log_final_document(debug_print, result)
                return result

        if detected_kind == "hwpx":
            result = {
                "url": url,
                "final_url": resp.url,
                "title": resolved_title,
                "content_type": "hwpx",
                "source_content_type": content_type,
                "content_disposition": content_disposition,
                "raw_html": "",
                "raw_text": "",
                "text": "",
                "blocks": [],
                "search_title": search_title,
                "search_snippet": search_snippet,
                "display_link": display_link,
                "source_engine": source_engine,
                "status_code": resp.status_code,
                "byte_size": byte_size,
                "warning": "HWPX detected but extractor is not implemented yet.",
            }
            _log_final_document(debug_print, result)
            return result

        result = {
            "url": url,
            "final_url": resp.url,
            "title": resolved_title,
            "content_type": "unknown",
            "source_content_type": content_type,
            "content_disposition": content_disposition,
            "raw_html": "",
            "raw_text": "",
            "text": "",
            "blocks": [],
            "search_title": search_title,
            "search_snippet": search_snippet,
            "display_link": display_link,
            "source_engine": source_engine,
            "status_code": resp.status_code,
            "byte_size": byte_size,
            "warning": f"unsupported or indeterminate content type: {content_type}",
        }
        _log_final_document(debug_print, result)
        return result

    except Exception as e:
        if debug_print:
            debug_print(f"EXTERNAL_FETCH ERROR | {url}", str(e))
        return {
            "url": url,
            "error": str(e),
            "search_title": search_title,
            "search_snippet": search_snippet,
            "display_link": display_link,
            "source_engine": source_engine,
        }


def run(
    query: str = "",
    urls: Optional[List[str]] = None,
    items: Optional[List[Dict[str, Any]]] = None,
    max_fetch: int = 5,
    timeout: int = DEFAULT_TIMEOUT,
    debug_print=None,
) -> ToolResult:
    try:
        normalized_inputs = _normalize_input_urls(urls=urls, items=items)

        if not normalized_inputs:
            return ToolResult(
                name="external_fetch",
                ok=False,
                data={"error": "no urls or items provided"},
            )

        max_fetch = max(1, int(max_fetch))
        selected_inputs = normalized_inputs[:max_fetch]

        if debug_print:
            debug_print(
                "EXTERNAL_FETCH START",
                {
                    "query": query,
                    "requested_count": len(normalized_inputs),
                    "selected_count": len(selected_inputs),
                    "urls": [x["url"] for x in selected_inputs],
                },
            )

        documents: List[Dict[str, Any]] = []

        for item in selected_inputs:
            doc = fetch_one_url(
                url=item["url"],
                search_title=item.get("search_title", ""),
                search_snippet=item.get("search_snippet", ""),
                display_link=item.get("display_link", ""),
                source_engine=item.get("source_engine", ""),
                timeout=timeout,
                debug_print=debug_print,
            )
            documents.append(doc)

        fetched_ok = [doc for doc in documents if not doc.get("error")]

        return ToolResult(
            name="external_fetch",
            ok=True,
            data={
                "query": query,
                "requested_count": len(normalized_inputs),
                "fetched_count": len(fetched_ok),
                "documents": documents,
            },
        )

    except Exception as e:
        if debug_print:
            debug_print("EXTERNAL_FETCH RUN ERROR", str(e))
        return ToolResult(
            name="external_fetch",
            ok=False,
            data={"error": str(e)},
        )