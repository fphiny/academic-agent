from __future__ import annotations

import json
import os
import re
import tempfile
import zipfile
from dataclasses import asdict, dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional
from urllib.parse import unquote, urlparse

import requests
from pypdf import PdfReader

from ..base import ToolResult, normalize_whitespace


DEFAULT_TIMEOUT = 20
MAX_RAW_TEXT_CHARS = 200_000

HTML_MIME_TYPES = {
    "text/html",
    "application/xhtml+xml",
}

TEXT_MIME_TYPES = {
    "text/plain",
    "text/markdown",
    "application/json",
    "text/csv",
    "application/xml",
    "text/xml",
}

PDF_MIME_TYPES = {
    "application/pdf",
}

DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

GENERIC_BINARY_MIME_TYPES = {
    "application/octet-stream",
    "application/x-msdownload",
    "application/download",
    "application/x-download",
    "binary/octet-stream",
}

HWP_MIME_TYPES = {
    "application/x-hwp",
    "application/haansofthwp",
    "application/hwp",
    "application/octet-stream",
    "application/x-msdownload",
}


@dataclass
class RawDocument:
    url: str
    final_url: str
    content_type: str
    status_code: int
    source_content_type: str = ""
    content_disposition: str = ""
    raw_html: str = ""
    raw_text: str = ""
    title: str = ""
    search_title: str = ""
    display_link: str = ""
    source_engine: str = ""
    warning: str = ""


def _safe_truncate(text: str, max_chars: int = MAX_RAW_TEXT_CHARS) -> str:
    text = normalize_whitespace(text or "")
    if len(text) <= max_chars:
        return text
    return text[:max_chars].rstrip() + " ..."


def _base_content_type(content_type: str) -> str:
    return (content_type or "").split(";")[0].strip().lower()


def _try_parse_json_like(value: Any) -> Any:
    if not isinstance(value, str):
        return value

    stripped = value.strip()
    if not stripped:
        return value

    if not (
        (stripped.startswith("[") and stripped.endswith("]"))
        or (stripped.startswith("{") and stripped.endswith("}"))
    ):
        return value

    try:
        return json.loads(stripped)
    except Exception:
        return value


def _extract_filename_from_url(url: str) -> str:
    try:
        path = unquote(urlparse(url).path)
        name = Path(path).name
        return name or ""
    except Exception:
        return ""


def _extract_filename_from_cd(content_disposition: str) -> str:
    if not content_disposition:
        return ""

    cd = content_disposition.strip()

    m = re.search(r"filename\*\s*=\s*UTF-8''([^;]+)", cd, re.I)
    if m:
        return unquote(m.group(1)).strip().strip('"')

    m = re.search(r'filename\s*=\s*"([^"]+)"', cd, re.I)
    if m:
        return m.group(1).strip()

    m = re.search(r"filename\s*=\s*([^;]+)", cd, re.I)
    if m:
        return m.group(1).strip().strip('"')

    return ""


def _file_suffix(filename: str) -> str:
    try:
        return Path(filename).suffix.lower()
    except Exception:
        return ""


def _resolved_filename(url: str, final_url: str, content_disposition: str) -> str:
    return (
        _extract_filename_from_cd(content_disposition)
        or _extract_filename_from_url(final_url)
        or _extract_filename_from_url(url)
    )


def _default_title(
    url: str,
    final_url: str = "",
    content_disposition: str = "",
    fallback: str = "",
) -> str:
    resolved_name = _resolved_filename(url, final_url, content_disposition)
    if resolved_name:
        return normalize_whitespace(resolved_name)
    return normalize_whitespace(fallback or "")


def _is_pdf_bytes(data: bytes) -> bool:
    return data.startswith(b"%PDF")


def _is_zip_bytes(data: bytes) -> bool:
    return data.startswith(b"PK\x03\x04")


def _is_ole_bytes(data: bytes) -> bool:
    return data.startswith(b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1")


def _looks_like_html_bytes(data: bytes) -> bool:
    if not data:
        return False

    head = data[:4096].lstrip()
    if not head:
        return False

    if b"\x00" in head[:512]:
        return False

    try:
        sample = head.decode("utf-8", errors="ignore").lower()
    except Exception:
        try:
            sample = head.decode("latin-1", errors="ignore").lower()
        except Exception:
            return False

    html_markers = (
        "<!doctype html",
        "<html",
        "<head",
        "<body",
        "<meta",
        "<title",
        "<script",
    )
    return any(marker in sample for marker in html_markers)


def _looks_like_text_bytes(data: bytes) -> bool:
    if not data:
        return False

    head = data[:4096]
    if not head:
        return False

    if b"\x00" in head:
        return False

    try:
        decoded = head.decode("utf-8", errors="ignore")
    except Exception:
        try:
            decoded = head.decode("latin-1", errors="ignore")
        except Exception:
            return False

    printable = sum(1 for ch in decoded if ch.isprintable() or ch in "\r\n\t")
    total = max(len(decoded), 1)
    return (printable / total) >= 0.9


def _head_hex(data: bytes, limit: int = 16) -> str:
    try:
        return data[:limit].hex()
    except Exception:
        return ""


def _detect_zip_kind(data: bytes, filename: str = "") -> str:
    if not _is_zip_bytes(data):
        return "unknown"

    try:
        with zipfile.ZipFile(BytesIO(data)) as zf:
            names = set(zf.namelist())

            if "[Content_Types].xml" in names and "word/document.xml" in names:
                return "docx"

            hwpx_markers = {
                "Contents/section0.xml",
                "Contents/content.hpf",
                "mimetype",
            }
            if any(name in names for name in hwpx_markers):
                if _file_suffix(filename) == ".hwpx":
                    return "hwpx"

            return "zip"
    except Exception:
        return "zip"


def _detect_ole_kind(data: bytes, filename: str = "") -> str:
    if not _is_ole_bytes(data):
        return "unknown"

    suffix = _file_suffix(filename)

    try:
        import olefile
    except Exception:
        if suffix == ".hwp":
            return "hwp"
        if suffix == ".doc":
            return "doc"
        return "ole"

    try:
        with olefile.OleFileIO(BytesIO(data)) as ole:
            stream_names = set("/".join(x) for x in ole.listdir())

            if "FileHeader" in stream_names:
                try:
                    raw = ole.openstream("FileHeader").read()
                    if b"HWP Document File" in raw:
                        return "hwp"
                except Exception:
                    pass
                return "hwp"

            if suffix == ".doc":
                return "doc"
            if suffix == ".hwp":
                return "hwp"

            return "ole"
    except Exception:
        if suffix == ".hwp":
            return "hwp"
        if suffix == ".doc":
            return "doc"
        return "ole"


def _guess_kind(
    *,
    url: str,
    final_url: str,
    content_type: str,
    content_disposition: str,
    content: bytes,
) -> str:
    ct = _base_content_type(content_type)
    preferred_filename = _resolved_filename(url, final_url, content_disposition)
    suffix = _file_suffix(preferred_filename)

    if _is_pdf_bytes(content):
        return "pdf"

    zip_kind = _detect_zip_kind(content, preferred_filename)
    if zip_kind == "docx":
        return "docx"
    if zip_kind == "hwpx":
        return "hwpx"

    ole_kind = _detect_ole_kind(content, preferred_filename)
    if ole_kind == "hwp":
        return "hwp"

    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".hwp":
        return "hwp"
    if suffix == ".hwpx":
        return "hwpx"
    if suffix in {".txt", ".md", ".csv", ".json", ".xml", ".log"}:
        return "text"
    if suffix in {".html", ".htm"}:
        return "html"

    if ct in PDF_MIME_TYPES:
        return "pdf"
    if ct in DOCX_MIME_TYPES:
        return "docx"
    if ct in HWP_MIME_TYPES:
        return "hwp"
    if ct in TEXT_MIME_TYPES:
        return "html" if _looks_like_html_bytes(content) else "text"
    if ct in HTML_MIME_TYPES:
        return "html"

    if _looks_like_html_bytes(content):
        return "html"
    if _looks_like_text_bytes(content):
        return "text"

    if ct in GENERIC_BINARY_MIME_TYPES:
        return "unknown"

    return "unknown"


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    texts: List[str] = []

    for page in reader.pages:
        try:
            extracted = page.extract_text() or ""
        except Exception:
            extracted = ""
        if extracted:
            texts.append(extracted)

    return _safe_truncate("\n".join(texts))


def _extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except Exception as e:
        raise RuntimeError("DOCX 지원을 위해 `python-docx` 설치가 필요합니다.") from e

    document = DocxDocument(BytesIO(data))
    texts: List[str] = []

    for para in document.paragraphs:
        text = normalize_whitespace(para.text or "")
        if text:
            texts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [normalize_whitespace(cell.text or "") for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                texts.append(" | ".join(cells))

    return _safe_truncate("\n\n".join(texts))


def _extract_hwp_text(data: bytes) -> str:
    try:
        from langchain_teddynote.document_loaders import HWPLoader
    except Exception as e:
        raise RuntimeError("HWP 지원을 위해 `langchain-teddynote` 설치가 필요합니다.") from e

    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".hwp", delete=False) as tmp:
            tmp.write(data)
            tmp.flush()
            tmp_path = tmp.name

        loader = HWPLoader(tmp_path)
        docs = loader.load()

        texts: List[str] = []
        for doc in docs:
            page_content = normalize_whitespace(getattr(doc, "page_content", "") or "")
            if page_content:
                texts.append(page_content)

        return _safe_truncate("\n\n".join(texts))
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass


def _decode_response_text(resp: requests.Response) -> str:
    try:
        return resp.text
    except Exception:
        try:
            encoding = resp.encoding or resp.apparent_encoding or "utf-8"
            return resp.content.decode(encoding, errors="replace")
        except Exception:
            return resp.content.decode("utf-8", errors="replace")


def _append_normalized_url(
    normalized: List[Dict[str, Any]],
    seen: set,
    link: Any,
    *,
    search_title: str = "",
    display_link: str = "",
    source_engine: str = "",
) -> None:
    link = _try_parse_json_like(link)

    if isinstance(link, (list, tuple)):
        for one_link in link:
            _append_normalized_url(
                normalized,
                seen,
                one_link,
                search_title=search_title,
                display_link=display_link,
                source_engine=source_engine,
            )
        return

    if isinstance(link, dict):
        _append_normalized_url(
            normalized,
            seen,
            link.get("link") or link.get("url"),
            search_title=str(link.get("title") or search_title or ""),
            display_link=str(link.get("displayLink") or display_link or ""),
            source_engine=str(link.get("source_engine") or source_engine or ""),
        )
        return

    link = str(link or "").strip()
    if not link or link in seen:
        return

    seen.add(link)
    normalized.append(
        {
            "url": link,
            "search_title": normalize_whitespace(search_title),
            "display_link": normalize_whitespace(display_link),
            "source_engine": normalize_whitespace(source_engine),
        }
    )


def _normalize_input_urls(
    url: Any = "",
    urls: Optional[Any] = None,
    items: Optional[List[Dict[str, Any]]] = None,
) -> List[Dict[str, Any]]:
    normalized: List[Dict[str, Any]] = []
    seen = set()

    url = _try_parse_json_like(url)
    urls = _try_parse_json_like(urls)
    items = _try_parse_json_like(items)

    if isinstance(url, dict):
        _append_normalized_url(
            normalized,
            seen,
            url.get("link") or url.get("url"),
            search_title=str(url.get("title") or ""),
            display_link=str(url.get("displayLink") or ""),
            source_engine=str(url.get("source_engine") or ""),
        )
    elif isinstance(url, (list, tuple)):
        merged_urls: List[Any] = list(url)
        if isinstance(urls, (list, tuple)):
            merged_urls.extend(list(urls))
        elif urls not in (None, "", []):
            merged_urls.append(urls)
        urls = merged_urls
    else:
        _append_normalized_url(normalized, seen, url)

    if isinstance(items, list):
        for item in items:
            item = _try_parse_json_like(item)
            if not isinstance(item, dict):
                continue
            _append_normalized_url(
                normalized,
                seen,
                item.get("link") or item.get("url"),
                search_title=str(item.get("title") or ""),
                display_link=str(item.get("displayLink") or ""),
                source_engine=str(item.get("source_engine") or ""),
            )

    if isinstance(urls, str):
        parsed_urls = _try_parse_json_like(urls)
        if isinstance(parsed_urls, (list, tuple, dict)):
            urls = parsed_urls
        else:
            _append_normalized_url(normalized, seen, urls)
    if isinstance(urls, dict):
        _append_normalized_url(
            normalized,
            seen,
            urls.get("link") or urls.get("url"),
            search_title=str(urls.get("title") or ""),
            display_link=str(urls.get("displayLink") or ""),
            source_engine=str(urls.get("source_engine") or ""),
        )
    elif isinstance(urls, (list, tuple)):
        for one_url in urls:
            one_url = _try_parse_json_like(one_url)
            if isinstance(one_url, dict):
                _append_normalized_url(
                    normalized,
                    seen,
                    one_url.get("link") or one_url.get("url"),
                    search_title=str(one_url.get("title") or ""),
                    display_link=str(one_url.get("displayLink") or ""),
                    source_engine=str(one_url.get("source_engine") or ""),
                )
            else:
                _append_normalized_url(normalized, seen, one_url)

    return normalized


def fetch_one_url(
    *,
    url: str,
    search_title: str = "",
    display_link: str = "",
    source_engine: str = "",
    timeout: int = DEFAULT_TIMEOUT,
    debug_print=None,
) -> Dict[str, Any]:
    try:
        resp = requests.get(
            url,
            timeout=timeout,
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/122.0.0.0 Safari/537.36"
                )
            },
            allow_redirects=True,
        )
        resp.raise_for_status()

        content_type = resp.headers.get("content-type", "").lower()
        content_disposition = resp.headers.get("content-disposition", "")
        final_url = resp.url
        detected_kind = _guess_kind(
            url=url,
            final_url=final_url,
            content_type=content_type,
            content_disposition=content_disposition,
            content=resp.content,
        )

        title = _default_title(
            url=url,
            final_url=final_url,
            content_disposition=content_disposition,
            fallback=search_title or url,
        )

        if debug_print:
            debug_print(
                f"FETCH_RAW RESPONSE META | {url}",
                {
                    "status_code": resp.status_code,
                    "content_type": content_type,
                    "content_disposition": content_disposition,
                    "final_url": final_url,
                    "detected_kind": detected_kind,
                    "title": title,
                    "head_hex": _head_hex(resp.content),
                    "byte_size": len(resp.content or b""),
                },
            )

        if detected_kind == "html":
            raw_html = _decode_response_text(resp)

            try:
                from bs4 import BeautifulSoup

                soup = BeautifulSoup(raw_html, "html.parser")
                if soup.title and soup.title.string:
                    title = normalize_whitespace(soup.title.string) or title
            except Exception:
                pass

            return asdict(
                RawDocument(
                    url=url,
                    final_url=final_url,
                    content_type="html",
                    source_content_type=content_type,
                    content_disposition=content_disposition,
                    status_code=resp.status_code,
                    raw_html=_safe_truncate(raw_html),
                    title=title,
                    search_title=search_title,
                    display_link=display_link,
                    source_engine=source_engine,
                )
            )

        if detected_kind == "text":
            text = _safe_truncate(_decode_response_text(resp))

            return asdict(
                RawDocument(
                    url=url,
                    final_url=final_url,
                    content_type="text",
                    source_content_type=content_type,
                    content_disposition=content_disposition,
                    status_code=resp.status_code,
                    raw_text=text,
                    title=title,
                    search_title=search_title,
                    display_link=display_link,
                    source_engine=source_engine,
                )
            )

        if detected_kind == "pdf":
            text = _extract_pdf_text(resp.content)

            return asdict(
                RawDocument(
                    url=url,
                    final_url=final_url,
                    content_type="pdf",
                    source_content_type=content_type,
                    content_disposition=content_disposition,
                    status_code=resp.status_code,
                    raw_text=text,
                    title=title,
                    search_title=search_title,
                    display_link=display_link,
                    source_engine=source_engine,
                )
            )

        if detected_kind == "docx":
            text = _extract_docx_text(resp.content)

            return asdict(
                RawDocument(
                    url=url,
                    final_url=final_url,
                    content_type="docx",
                    source_content_type=content_type,
                    content_disposition=content_disposition,
                    status_code=resp.status_code,
                    raw_text=text,
                    title=title,
                    search_title=search_title,
                    display_link=display_link,
                    source_engine=source_engine,
                )
            )

        if detected_kind == "hwp":
            try:
                text = _extract_hwp_text(resp.content)
                return asdict(
                    RawDocument(
                        url=url,
                        final_url=final_url,
                        content_type="hwp",
                        source_content_type=content_type,
                        content_disposition=content_disposition,
                        status_code=resp.status_code,
                        raw_text=text,
                        title=title,
                        search_title=search_title,
                        display_link=display_link,
                        source_engine=source_engine,
                    )
                )
            except Exception as e:
                return asdict(
                    RawDocument(
                        url=url,
                        final_url=final_url,
                        content_type="hwp",
                        source_content_type=content_type,
                        content_disposition=content_disposition,
                        status_code=resp.status_code,
                        raw_text="",
                        title=title,
                        search_title=search_title,
                        display_link=display_link,
                        source_engine=source_engine,
                        warning=f"hwp extraction failed: {e}",
                    )
                )

        if detected_kind == "hwpx":
            return asdict(
                RawDocument(
                    url=url,
                    final_url=final_url,
                    content_type="hwpx",
                    source_content_type=content_type,
                    content_disposition=content_disposition,
                    status_code=resp.status_code,
                    raw_text="",
                    title=title,
                    search_title=search_title,
                    display_link=display_link,
                    source_engine=source_engine,
                    warning="hwpx detected but extractor is not implemented yet",
                )
            )

        return asdict(
            RawDocument(
                url=url,
                final_url=final_url,
                content_type="unknown",
                source_content_type=content_type,
                content_disposition=content_disposition,
                status_code=resp.status_code,
                raw_html="",
                raw_text="",
                title=title,
                search_title=search_title,
                display_link=display_link,
                source_engine=source_engine,
                warning=f"unsupported or indeterminate content type: {content_type}",
            )
        )

    except Exception as e:
        if debug_print:
            debug_print(f"FETCH_RAW ERROR | {url}", str(e))
        return {
            "url": url,
            "final_url": url,
            "content_type": "",
            "source_content_type": "",
            "content_disposition": "",
            "status_code": 0,
            "raw_html": "",
            "raw_text": "",
            "title": "",
            "error": str(e),
            "search_title": search_title,
            "display_link": display_link,
            "source_engine": source_engine,
        }


def run(
    query: str = "",
    url: Any = "",
    urls: Optional[Any] = None,
    items: Optional[List[Dict[str, Any]]] = None,
    max_fetch: int = 5,
    timeout: int = DEFAULT_TIMEOUT,
    debug_print=None,
) -> ToolResult:
    try:
        normalized_inputs = _normalize_input_urls(
            url=url,
            urls=urls,
            items=items,
        )

        if debug_print:
            debug_print(
                "FETCH_RAW NORMALIZED INPUTS",
                {
                    "query": query,
                    "raw_url": url,
                    "raw_urls": urls,
                    "raw_items_count": len(items) if isinstance(items, list) else 0,
                    "normalized_count": len(normalized_inputs),
                    "normalized_urls": [x["url"] for x in normalized_inputs],
                },
                max_len=12000,
            )

        if not normalized_inputs:
            return ToolResult(
                name="external_fetch_raw",
                ok=False,
                data={
                    "error": "no url, urls, or items provided",
                    "received": {
                        "url": url,
                        "urls": urls,
                        "items_count": len(items) if isinstance(items, list) else 0,
                    },
                },
            )

        max_fetch = max(1, int(max_fetch))
        selected_inputs = normalized_inputs[:max_fetch]

        if debug_print:
            debug_print(
                "FETCH_RAW START",
                {
                    "query": query,
                    "requested_count": len(normalized_inputs),
                    "selected_count": len(selected_inputs),
                    "urls": [x["url"] for x in selected_inputs],
                },
            )

        documents: List[Dict[str, Any]] = []

        for item in selected_inputs:
            doc = fetch_one_url(
                url=item["url"],
                search_title=item.get("search_title", ""),
                display_link=item.get("display_link", ""),
                source_engine=item.get("source_engine", ""),
                timeout=timeout,
                debug_print=debug_print,
            )
            documents.append(doc)

        fetched_ok = [doc for doc in documents if not doc.get("error")]

        return ToolResult(
            name="external_fetch_raw",
            ok=True,
            data={
                "query": query,
                "requested_count": len(normalized_inputs),
                "fetched_count": len(fetched_ok),
                "documents": documents,
            },
        )

    except Exception as e:
        if debug_print:
            debug_print("FETCH_RAW RUN ERROR", str(e))
        return ToolResult(
            name="external_fetch_raw",
            ok=False,
            data={"error": str(e)},
        )